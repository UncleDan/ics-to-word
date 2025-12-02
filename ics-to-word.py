import sys
import os
from datetime import datetime, timezone, timedelta
from icalendar import Calendar
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from typing import Optional

class ICSConverterGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Convertitore ICS to Word")
        self.root.geometry("500x300")
        
        # Imposta l'icona (se disponibile)
        try:
            self.root.iconbitmap(default='icon.ico')
        except:
            pass
        
        self.ics_file_path = ""
        
        # Configura lo stile
        self.setup_styles()
        
        # Crea l'interfaccia
        self.create_widgets()
        
    def setup_styles(self):
        """Configura gli stili per l'interfaccia"""
        style = ttk.Style()
        style.configure("Title.TLabel", font=("Arial", 14, "bold"))
        style.configure("Status.TLabel", font=("Arial", 10))
        
    def create_widgets(self):
        """Crea i widget dell'interfaccia grafica"""
        
        # Frame principale
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Titolo
        title_label = ttk.Label(main_frame, text="Convertitore Calendario ICS to Word", 
                               style="Title.TLabel")
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        # Etichetta file selezionato
        self.file_label = ttk.Label(main_frame, text="Nessun file selezionato", 
                                   wraplength=400)
        self.file_label.grid(row=1, column=0, columnspan=3, pady=(0, 20))
        
        # Pulsante per selezionare file
        select_button = ttk.Button(main_frame, text="Seleziona file ICS", 
                                  command=self.select_ics_file)
        select_button.grid(row=2, column=0, padx=5, pady=10)
        
        # Pulsante per convertire
        self.convert_button = ttk.Button(main_frame, text="Converti in Word", 
                                        command=self.convert_file,
                                        state="disabled")
        self.convert_button.grid(row=2, column=1, padx=5, pady=10)
        
        # Pulsante per uscire
        exit_button = ttk.Button(main_frame, text="Esci", 
                                command=self.root.quit)
        exit_button.grid(row=2, column=2, padx=5, pady=10)
        
        # Barra di progresso
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate', length=300)
        self.progress.grid(row=3, column=0, columnspan=3, pady=(20, 10))
        
        # Etichetta di stato
        self.status_label = ttk.Label(main_frame, text="Pronto", 
                                     style="Status.TLabel")
        self.status_label.grid(row=4, column=0, columnspan=3)
        
        # Configura il resize
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        
    def select_ics_file(self):
        """Apre la finestra di dialogo per selezionare un file ICS"""
        file_path = filedialog.askopenfilename(
            title="Seleziona file ICS",
            filetypes=[
                ("File ICS", "*.ics"),
                ("File iCalendar", "*.ical"),
                ("Tutti i file", "*.*")
            ]
        )
        
        if file_path:
            self.ics_file_path = file_path
            self.file_label.config(text=f"File selezionato:\n{file_path}")
            self.convert_button.config(state="normal")
            self.status_label.config(text="File selezionato. Pronto per la conversione.")
    
    def convert_file(self):
        """Avvia la conversione del file"""
        if not self.ics_file_path:
            messagebox.showerror("Errore", "Seleziona prima un file ICS!")
            return
        
        # Disabilita i pulsanti durante la conversione
        self.convert_button.config(state="disabled")
        self.progress.start()
        self.status_label.config(text="Conversione in corso...")
        self.root.update()
        
        try:
            # Genera il nome del file di output
            base_name = os.path.splitext(os.path.basename(self.ics_file_path))[0]
            output_dir = os.path.dirname(self.ics_file_path)
            docx_file_path = os.path.join(output_dir, f"{base_name}.Calendar.docx")
            
            # Controlla se il file di output esiste già
            if os.path.exists(docx_file_path):
                response = messagebox.askyesno(
                    "File esistente",
                    f"Il file '{os.path.basename(docx_file_path)}' esiste già.\nSovrascriverlo?"
                )
                if not response:
                    # Chiede all'utente di specificare un nuovo nome
                    new_path = filedialog.asksaveasfilename(
                        title="Salva come",
                        defaultextension=".docx",
                        initialfile=f"{base_name}.Calendar.docx",
                        filetypes=[("Documenti Word", "*.docx"), ("Tutti i file", "*.*")]
                    )
                    if new_path:
                        docx_file_path = new_path
                    else:
                        self.reset_ui()
                        return
            
            # Esegui la conversione
            success = convert_ics_to_word(self.ics_file_path, docx_file_path)
            
            if success:
                messagebox.showinfo(
                    "Successo", 
                    f"Conversione completata!\n\n"
                    f"File creato:\n{docx_file_path}"
                )
                self.status_label.config(text="Conversione completata con successo!")
            else:
                messagebox.showerror(
                    "Errore", 
                    "Si è verificato un errore durante la conversione.\n"
                    "Controlla il formato del file ICS e riprova."
                )
                self.status_label.config(text="Errore durante la conversione")
                
        except Exception as e:
            messagebox.showerror("Errore", f"Errore imprevisto: {str(e)}")
            self.status_label.config(text=f"Errore: {str(e)}")
        finally:
            # Riabilita l'interfaccia
            self.progress.stop()
            self.convert_button.config(state="normal")
    
    def reset_ui(self):
        """Ripristina l'interfaccia allo stato iniziale"""
        self.progress.stop()
        self.convert_button.config(state="disabled")
        self.status_label.config(text="Pronto")

def normalize_datetime(dt):
    """
    Normalizza un datetime per il confronto.
    Se è offset-aware, lo converte in offset-naive nel fuso orario locale.
    
    Args:
        dt: datetime object (può essere offset-naive o offset-aware)
        
    Returns:
        datetime: datetime offset-naive normalizzato
    """
    if dt is None:
        return None
    
    # Se è offset-aware, converti in offset-naive
    if hasattr(dt, 'tzinfo') and dt.tzinfo is not None:
        # Rimuovi le informazioni sul fuso orario
        return dt.replace(tzinfo=None)
    return dt

def convert_ics_to_word(ics_file_path, docx_file_path):
    """
    Converte un file ICS in un documento Word formattato.
    
    Args:
        ics_file_path (str): Percorso del file .ics
        docx_file_path (str): Percorso del file .docx di output
        
    Returns:
        bool: True se la conversione è riuscita, False altrimenti
    """
    
    # Leggi il file ICS
    try:
        with open(ics_file_path, 'r', encoding='utf-8') as f:
            ics_content = f.read()
    except UnicodeDecodeError:
        try:
            with open(ics_file_path, 'r', encoding='latin-1') as f:
                ics_content = f.read()
        except Exception as e:
            print(f"Errore nella lettura del file: {e}")
            return False
    except Exception as e:
        print(f"Errore nell'apertura del file: {e}")
        return False
    
    # Parsing del calendario ICS
    try:
        cal = Calendar.from_ical(ics_content)
    except Exception as e:
        print(f"Errore nel parsing del file ICS: {e}")
        return False
    
    # Crea un nuovo documento Word
    doc = Document()
    
    # Configurazione pagina
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Crea stili personalizzati
    # Stile per il titolo principale
    title_style = doc.styles.add_style('CalendarTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 32, 96)  # Blu scuro
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Stile per la data dell'evento
    date_style = doc.styles.add_style('EventDate', WD_STYLE_TYPE.PARAGRAPH)
    date_style.font.name = 'Calibri'
    date_style.font.size = Pt(14)
    date_style.font.bold = True
    date_style.font.color.rgb = RGBColor(46, 117, 182)  # Blu
    date_style.paragraph_format.space_before = Pt(12)
    date_style.paragraph_format.space_after = Pt(6)
    
    # Stile per il titolo dell'evento
    event_title_style = doc.styles.add_style('EventTitle', WD_STYLE_TYPE.PARAGRAPH)
    event_title_style.font.name = 'Calibri'
    event_title_style.font.size = Pt(12)
    event_title_style.font.bold = True
    event_title_style.paragraph_format.space_after = Pt(3)
    
    # Stile per i dettagli dell'evento
    detail_style = doc.styles.add_style('EventDetail', WD_STYLE_TYPE.PARAGRAPH)
    detail_style.font.name = 'Calibri'
    detail_style.font.size = Pt(11)
    detail_style.paragraph_format.left_indent = Inches(0.25)
    detail_style.paragraph_format.space_after = Pt(3)
    
    # Titolo del documento con il nome del file
    base_name = os.path.splitext(os.path.basename(ics_file_path))[0]
    title = doc.add_paragraph()
    title.add_run(f'CALENDARIO - {base_name.upper()}')
    title.style = title_style
    
    # Aggiungi data di generazione
    gen_date = doc.add_paragraph()
    gen_date.add_run(f'Generato il: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_date.style = detail_style
    
    doc.add_paragraph()  # Spazio vuoto
    
    # Estrai e organizza gli eventi
    events = []
    
    for component in cal.walk():
        if component.name == "VEVENT":
            event = {}
            
            # Estrai titolo
            if 'summary' in component:
                event['title'] = str(component.get('summary')).strip()
            else:
                event['title'] = 'Senza titolo'
            
            # Estrai descrizione
            if 'description' in component:
                description = str(component.get('description')).strip()
                # Pulisci la descrizione da caratteri speciali multipli
                description = ' '.join(description.split())
                event['description'] = description
            else:
                event['description'] = ''
            
            # Estrai data di inizio
            if 'dtstart' in component:
                dt_start = component.get('dtstart').dt
                if isinstance(dt_start, datetime):
                    # Normalizza il datetime per il confronto
                    event['start'] = normalize_datetime(dt_start)
                    event['date_str'] = dt_start.strftime('%d/%m/%Y')
                    event['time_str'] = dt_start.strftime('%H:%M')
                else:
                    # È una data senza ora
                    event['start'] = datetime.combine(dt_start, datetime.min.time())
                    event['date_str'] = dt_start.strftime('%d/%m/%Y')
                    event['time_str'] = 'Tutto il giorno'
            else:
                event['start'] = None
                event['date_str'] = 'Data non specificata'
                event['time_str'] = ''
            
            # Estrai data di fine
            if 'dtend' in component:
                dt_end = component.get('dtend').dt
                if isinstance(dt_end, datetime):
                    # Normalizza il datetime per il confronto
                    event['end'] = normalize_datetime(dt_end)
                    event['end_time_str'] = dt_end.strftime('%H:%M')
                else:
                    # È una data senza ora
                    event['end'] = datetime.combine(dt_end, datetime.min.time())
                    event['end_time_str'] = 'Tutto il giorno'
            else:
                event['end'] = None
                event['end_time_str'] = ''
            
            # Estrai luogo
            if 'location' in component:
                location = str(component.get('location')).strip()
                event['location'] = location
            else:
                event['location'] = ''
            
            # Estrai UID (identificativo unico)
            if 'uid' in component:
                event['uid'] = str(component.get('uid'))
            else:
                event['uid'] = ''
            
            events.append(event)
    
    # Ordina eventi per data
    # Usa una funzione di chiave personalizzata che gestisce i None
    def get_sort_key(event):
        if event['start'] is None:
            # Restituisci una data molto lontana nel passato per gli eventi senza data
            return datetime.min
        return event['start']
    
    events.sort(key=get_sort_key)
    
    # Aggiungi eventi al documento Word
    current_date = None
    event_count = 0
    
    for event in events:
        # Se la data è cambiata, aggiungi un'intestazione di data
        if event['date_str'] != current_date:
            current_date = event['date_str']
            
            # Linea separatrice
            doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data
            date_para = doc.add_paragraph()
            date_para.add_run(event['date_str'].upper())
            date_para.style = date_style
        
        # Titolo evento
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"• {event['title']}")
        title_para.style = event_title_style
        
        # Ora
        if event['time_str']:
            time_para = doc.add_paragraph()
            time_text = f"🕒 {event['time_str']}"
            if event['end_time_str'] and event['end_time_str'] != event['time_str']:
                time_text += f" - {event['end_time_str']}"
            time_run = time_para.add_run(time_text)
            time_para.style = detail_style
        
        # Luogo
        if event['location']:
            location_para = doc.add_paragraph()
            location_para.add_run(f"📍 {event['location']}")
            location_para.style = detail_style
        
        # Descrizione
        if event['description']:
            desc_para = doc.add_paragraph()
            desc_para.add_run(f"📝 {event['description']}")
            desc_para.style = detail_style
        
        # Spazio tra eventi
        doc.add_paragraph().style = detail_style
        event_count += 1
    
    # Riepilogo
    if events:
        doc.add_page_break()
        
        summary_title = doc.add_paragraph()
        summary_title.add_run('RIEPILOGO EVENTI')
        summary_title.style = title_style
        
        summary_info = doc.add_paragraph()
        summary_info.add_run(f'Totale eventi: {event_count}')
        summary_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        summary_info.style = detail_style
        
        doc.add_paragraph()  # Spazio
        
        summary_table = doc.add_table(rows=1, cols=4)
        summary_table.style = 'Light Grid Accent 1'
        
        # Intestazioni tabella
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = 'Data'
        header_cells[1].text = 'Ora'
        header_cells[2].text = 'Evento'
        header_cells[3].text = 'Luogo'
        
        # Applica grassetto alle intestazioni
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Aggiungi eventi alla tabella riepilogativa
        for event in events:
            row_cells = summary_table.add_row().cells
            row_cells[0].text = event['date_str']
            row_cells[1].text = event['time_str']
            row_cells[2].text = event['title']
            row_cells[3].text = event['location']
    
    # Salva il documento
    try:
        doc.save(docx_file_path)
        print(f"Documento Word creato con successo: {docx_file_path}")
        print(f"Totale eventi elaborati: {event_count}")
        return True
    except Exception as e:
        print(f"Errore nel salvataggio del documento: {e}")
        return False

def check_dependencies():
    """Controlla e installa le dipendenze mancanti"""
    dependencies = [
        ('icalendar', 'icalendar'),
        ('docx', 'python-docx'),
        ('tkinter', 'tkinter')  # tkinter è di solito incluso con Python
    ]
    
    for import_name, package_name in dependencies:
        try:
            if import_name == 'icalendar':
                from icalendar import Calendar
            elif import_name == 'docx':
                from docx import Document
            elif import_name == 'tkinter':
                import tkinter
            print(f"{package_name}: OK")
        except ImportError:
            print(f"Installazione di {package_name}...")
            try:
                import subprocess
                subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])
                print(f"{package_name} installato con successo")
            except Exception as e:
                print(f"Errore nell'installazione di {package_name}: {e}")
                return False
    return True

def main():
    """Funzione principale"""
    print("=" * 50)
    print("Convertitore ICS to Word con interfaccia grafica")
    print("=" * 50)
    
    # Controlla dipendenze
    if not check_dependencies():
        print("Impossibile installare tutte le dipendenze.")
        response = input("Vuoi provare comunque? (s/n): ")
        if response.lower() != 's':
            return
    
    try:
        # Crea l'interfaccia grafica
        root = tk.Tk()
        app = ICSConverterGUI(root)
        
        # Imposta l'icona della finestra (Windows)
        try:
            root.iconbitmap('icon.ico')
        except:
            pass
        
        # Avvia l'applicazione
        root.mainloop()
        
    except Exception as e:
        print(f"Errore nell'avvio dell'interfaccia grafica: {e}")
        
        # Fallback: interfaccia da riga di comando
        print("\nAvvio interfaccia da riga di comando...")
        if len(sys.argv) == 3:
            ics_file = sys.argv[1]
            docx_file = sys.argv[2]
            convert_ics_to_word(ics_file, docx_file)
        else:
            print("Utilizzo: python ics_to_word_gui.py <file_input.ics> <file_output.docx>")
            print("Esempio: python ics_to_word_gui.py calendario.ics output.docx")

if __name__ == "__main__":
    main()